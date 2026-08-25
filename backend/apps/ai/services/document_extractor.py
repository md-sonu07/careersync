import re
import io
import logging

logger = logging.getLogger(__name__)


class DocumentExtractionError(Exception):
    """Raised when text extraction fails or file is corrupt."""
    pass


class DocumentExtractor:
    """
    Modular Document Extraction Service supporting PDF, DOCX, and TXT.
    Includes text cleaning and OCR/scanned document detection.
    """

    MAX_TEXT_LENGTH = 150000  # Cap extracted text length to prevent memory/token overflows

    @classmethod
    def extract_text(cls, file_obj, filename, file_type):
        """
        Main entry point for extracting text from a file object.
        Returns a tuple: (extracted_text: str, ocr_required: bool)
        """
        file_ext = (file_type or filename.split('.')[-1]).lower().strip('.')

        if file_ext in ['pdf', 'application/pdf']:
            text, ocr = cls._extract_from_pdf(file_obj)
        elif file_ext in ['docx', 'doc', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            text, ocr = cls._extract_from_docx(file_obj)
        elif file_ext in ['txt', 'text/plain']:
            text, ocr = cls._extract_from_txt(file_obj)
        else:
            raise DocumentExtractionError(f"Unsupported file format: '.{file_ext}'")

        cleaned_text = cls._clean_text(text)

        # Detect scanned/empty text (< 5 words)
        word_count = len(cleaned_text.split())
        if word_count < 5 and not ocr:
            ocr = True

        return cleaned_text[:cls.MAX_TEXT_LENGTH], ocr

    @classmethod
    def _extract_from_pdf(cls, file_obj):
        """Extract text from PDF using PyMuPDF (fitz) with pypdf fallback."""
        text = ""
        ocr_required = False

        # Attempt 1: PyMuPDF (fitz)
        try:
            import pymupdf
            file_bytes = file_obj.read() if hasattr(file_obj, 'read') else file_obj
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            full_text = []
            image_count = 0

            for page in doc:
                page_text = page.get_text("text") or ""
                full_text.append(page_text)
                image_count += len(page.get_images())

            text = "\n\n".join(full_text).strip()

            # If text is extremely short but page has images, OCR is required
            if len(text.split()) < 20 and image_count > 0:
                ocr_required = True

            if text.strip():
                return text, ocr_required
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed or missing: {e}. Trying pypdf...")

        # Attempt 2: pypdf fallback
        try:
            import pypdf
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
            reader = pypdf.PdfReader(file_obj if hasattr(file_obj, 'read') else io.BytesIO(file_obj))
            full_text = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                full_text.append(page_text)

            text = "\n\n".join(full_text).strip()
            if len(text.split()) < 20:
                ocr_required = True
            return text, ocr_required
        except Exception as e:
            logger.error(f"pypdf extraction also failed: {e}")
            raise DocumentExtractionError(f"Could not extract text from PDF: {str(e)}")

    @classmethod
    def _extract_from_docx(cls, file_obj):
        """Extract text from DOCX using python-docx."""
        try:
            import docx
            file_bytes = file_obj.read() if hasattr(file_obj, 'read') else file_obj
            doc = docx.Document(io.BytesIO(file_bytes))

            full_text = []
            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    full_text.append(p.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            text = "\n\n".join(full_text).strip()
            return text, False
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise DocumentExtractionError(f"Could not extract text from DOCX file: {str(e)}")

    @classmethod
    def _extract_from_txt(cls, file_obj):
        """Extract text from TXT file with safe encoding fallbacks."""
        try:
            content_bytes = file_obj.read() if hasattr(file_obj, 'read') else file_obj

            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
            text = None
            for enc in encodings:
                try:
                    text = content_bytes.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                text = content_bytes.decode('utf-8', errors='ignore')

            return text.strip(), False
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise DocumentExtractionError(f"Could not read text file: {str(e)}")

    @classmethod
    def _clean_text(cls, text):
        """Clean whitespace, remove control characters, and normalize headers."""
        if not text:
            return ""

        # Replace carriage returns & multiple newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Replace multiple spaces/tabs with single space per line
        lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
        # Collapse >2 consecutive empty lines
        cleaned = re.sub(r'\n{3,}', '\n\n', '\n'.join(lines))
        return cleaned.strip()
